#!/usr/bin/env python
# coding: utf-8


import logging
from datetime import datetime

LOG_PATH = "sms_qbo_debug.log"

logging.basicConfig(
    filename=LOG_PATH,
    level=logging.ERROR,
    format="%(asctime)s %(levelname)s %(message)s",
)

def get_intuit_tid(response):
    return response.headers.get("intuit_tid") or response.headers.get("Intuit-Tid")

def log_qbo_error(context, response, extra=None):
    """
    Logs key QBO error info including intuit_tid to share logs with Intuit support.
    """
    timestamp = datetime.now().isoformat(timespec="seconds")
    tid = get_intuit_tid(response)

    msg = {
        "timestamp": timestamp,
        "context": context,
        "status_code": response.status_code,
        "intuit_tid": tid,
        "response_text": response.text,
        "extra": extra,
    }

    logging.error("QBO_ERROR %s", msg)
    print(
        f"❌ [{timestamp}] QBO error in {context} | "
        f"status={response.status_code} | intuit_tid={tid}"
    )


# In[6]:


#pwd
import math


# In[7]:


from dotenv import load_dotenv
import os

load_dotenv()

ACCESS_TOKEN = os.getenv("ACCESS_TOKEN")
REALM_ID = os.getenv("REALM_ID")

if not REALM_ID:
    raise RuntimeError("Missing REALM_ID in environment (.env).")

QBO_ENV = os.getenv("QBO_ENV", "sandbox").lower()
BASE_URL = (
    "https://sandbox-quickbooks.api.intuit.com"
    if QBO_ENV == "sandbox"
    else "https://quickbooks.api.intuit.com"
)

# print(f"ACCESS_TOKEN: {ACCESS_TOKEN[:10]}...")  # Optional sanity check


# In[8]:


##### Init Results function

def send_to_results(d):
    global results_box, results_items

    if " x " in d["size"]:
        size_formatted = f"{d['size'].split(' x ')[0]}\" x {d['size'].split(' x ')[1]}\""
    else:
        size_formatted = ""

    block_text = f"{d['print_type']}\n"
    if d["canvas_cost"] is not None:
        block_text += f"    Regular Print Price:{'.' * (21 - len('    Print:'))} $ {d['canvas_cost']:.2f}\n"
    if d["pro_canvas_cost"] is not None:
        block_text += f"    Professional Print Price:{'.' * (31 - len('    Print (Pro Discount):'))} $ {d['pro_canvas_cost']:.2f}\n"
    if d["frame_cost"] > 0:
        block_text += f"    Frame Wood:{'.' * (35 - len('    Frame Wood:'))} $ {d['frame_cost']:.2f}\n"
    if d["stretching_fee"] > 0:
        block_text += f"    Stretch:{'.' * (35 - len('    Stretch:'))} $ {d['stretching_fee']:.2f}\n"
    if d["bracer_cost"] > 0:
        block_text += f"    Bracer Wood:{'.' * (35 - len('    Bracer Wood:'))} $ {d['bracer_cost']:.2f}\n"
    if d["upcharge"] > 0:
        block_text += f"    ≥ 72\" Upcharge:{'.' * (35 - len('    ≥ 72\" Upcharge:'))} $ {d['upcharge']:.2f}\n"
    
    block_text += "     " + ("⋯⋯⋯⋯⋯" * 3) + "\n"
    
    # Compute total costs for display
    if d["canvas_cost"] is not None:
        regular_total = d["canvas_cost"] + d["frame_cost"] + d["stretching_fee"] + d["bracer_cost"] + d["upcharge"]
        block_text += f"    Regular Total:{'.' * (35 - len('    Regular Total:'))} $ {regular_total:.2f}\n"
    
    if d["pro_canvas_cost"] is not None:
        pro_total = d["pro_canvas_cost"] + d["frame_cost"] + d["stretching_fee"] + d["bracer_cost"] + d["upcharge"]
        block_text += f"    Professional Total:{'.' * (35 - len('    Professional Total:'))} $ {pro_total:.2f}\n"


    divider = "=" * 45
    block_lines = block_text.strip().splitlines()

    results_box.configure(state="normal")
    
    # Start line index
    start_index = results_box.index("end-1c").split(".")[0]
    start_line = int(start_index)
    
    # Define color BEFORE using it
    colors = ["#ccff00", "#00ccff", "#ff00cc", "#ffcc00", "#cc00ff", "#00ffcc", "#ff0000"]
    tag_color = colors[len(results_items) % len(colors)]
    
    # Insert main text block
    results_box.insert("end", block_text + "\n")
    
    # Center tag for the button
    results_box.tag_configure("center", justify="center")

    try:
        current_scroll = results_box.yview()
        top_line = int(results_box.index("@0,0").split('.')[0])
        bottom_line = int(results_box.index(f"@0,{results_box.winfo_height()}").split('.')[0])
    except Exception:
        current_scroll = (0.0, 1.0)
        top_line, bottom_line = 0, 9999

    
    # Create centered frame for button
    send_button = tk.Button(
        results_box,
        text="✔️",
        font=("Avenir Next", 10),
        command=lambda d=d: send_to_draft({**d, "from_results": True}),
        bg="#222222", fg="white", activebackground="#333333",
        bd=0, highlightthickness=0, padx=2, pady=0
    )
    results_box.insert(tk.END, "  ")
    results_box.window_create("end", window=send_button)
    results_box.insert("end", "\n")


    # Check if selected results block is still in view — if not, scroll to it
    if hasattr(results_box, "focused_result_line"):  # optional marker
        focus_line = results_box.focused_result_line
        if not (top_line <= focus_line <= bottom_line):
            results_box.see(f"{focus_line}.0")
    else:
        results_box.yview_moveto(current_scroll[0])
    
    # Insert divider and tag it
    results_box.insert("end", divider + "\n\n")
    divider_index = results_box.index("end-3l").split(".")[0]
    results_box.tag_configure(f"divider{len(results_items)}", font=("Monaco", 15, "bold"), foreground=tag_color)
    results_box.tag_add(f"divider{len(results_items)}", f"{divider_index}.0", f"{divider_index}.end")
    
    # Tag the title line
    results_box.tag_configure(f"title{len(results_items)}", font=("Monaco", 15, "bold"), foreground=tag_color)
    results_box.tag_add(f"title{len(results_items)}", f"{start_line}.0", f"{start_line}.end")

    results_box.configure(state="disabled")
    results_items.append(d)


# In[9]:


### Def helper functions to manage the custom item list of dictionaries

def remove_custom_item(index):
    if current_title in custom_items_by_title:
        del custom_items_by_title[current_title][index]
        update_draft_display()

def clear_custom_items():
    if current_title in custom_items_by_title:
        custom_items_by_title[current_title].clear()
        update_draft_display()

import uuid 

def add_custom_item():
    name = custom_item_name_entry.get().strip()
    desc = custom_item_desc_entry.get().strip()
    qty = custom_item_qty_var.get()
    unit_price = custom_item_price_var.get()

    if not name or qty <= 0 or unit_price <= 0:
        messagebox.showwarning("Missing Info", "Please enter a name, quantity > 0, and unit price > 0.")
        return

    item = {
        "id": str(uuid.uuid4()),          # ✅ add this
        "name": name,
        "description": desc,
        "quantity": qty,
        "unit_price": unit_price
    }
    if current_title:
        custom_items_by_title.setdefault(current_title, []).append(item)
    else:
        messagebox.showwarning("Missing Title", "Please select or enter a title before adding a custom item.")
        return


    # Clear custom item inputs after adding
    custom_item_name_entry.delete(0, tk.END)
    custom_item_desc_entry.delete(0, tk.END)
    custom_item_qty_var.set(0)
    custom_item_price_var.set(0.0)

    update_draft_display()

def add_custom_item_to_draft():
    add_custom_item()

    if current_title and custom_items_by_title.get(current_title):
        item = custom_items_by_title[current_title][-1]

        artist_first = artist_first_entry.get().strip()
        artist_last  = artist_last_entry.get().strip()

        draft_items.append({
            "linked_title": current_title,
            "print_type": item["name"],
            "size": item["description"],
            "num_prints": item["quantity"],
            "regular_price": item["unit_price"],
            "pro_price": item["unit_price"],
            "color": "#E9967A",
            "source": "custom",
            "custom_id": item["id"],       # ✅ add this
            "artist_first": artist_first,
            "artist_last": artist_last,
        })

    update_draft_display()


# In[10]:


def colorize_results():
    global color_index, colors
    color_index = 0  # reset for consistency

    line_count = int(results_box.index("end-1c").split('.')[0])
    current_color = colors[color_index % len(colors)]

    for i in range(1, line_count + 1):
        line_text = results_box.get(f"{i}.0", f"{i}.end")

        # Detect print type title line
        if line_text.startswith("Canvas with") or \
           line_text.startswith("Photorag") or \
           line_text.startswith("Enhanced Matte") or \
           line_text.startswith("Watercolor") or \
           line_text.startswith("Unstretched Canvas"):

            current_color = colors[color_index % len(colors)]
            results_box.tag_configure(f"title_{color_index}", font=("Monaco", 15, "bold"), foreground=current_color)
            results_box.tag_add(f"title_{color_index}", f"{i}.0", f"{i}.end")

        # Detect separator and highlight totals above
        elif line_text.strip() == "=" * 45:
            results_box.tag_configure(f"separator_{color_index}", font=("Monaco", 15, "bold"), foreground=current_color)
            results_box.tag_add(f"separator_{color_index}", f"{i}.0", f"{i}.end")

            # Highlight the two lines above the separator
            results_box.tag_configure(f"highlight_{color_index}", font=("Monaco", 12, "bold"), foreground="#ffffff")
            results_box.tag_add(f"highlight_{color_index}", f"{i-2}.0", f"{i}.0")

            color_index += 1  # advance for next block


# In[11]:


##### Function to send from input to draft
def send_to_draft(d):
    global draft_items, current_artist, current_title

    current_artist = artist_first_entry.get().strip() + " " + artist_last_entry.get().strip()
    current_title = title_entry.get().strip()
    if current_title and current_title not in draft_titles:
        draft_titles.append(current_title)
    ##print(f"📥 send_to_draft() called with current_title: '{current_title}'")

    # ⏬ Save scroll position
    try:
        previous_scroll = results_box.yview()
    except:
        previous_scroll = None


    if " x " in d["size"]:
        size_formatted = f"{d['size'].split(' x ')[0]}\" x {d['size'].split(' x ')[1]}\""
    else:
        size_formatted = ""

    
    item = {
        "print_type": d["print_type"],
        "size": size_formatted,
        "canvas_cost": d["canvas_cost"] or 0.00,
        "pro_canvas_cost": d["pro_canvas_cost"] or 0.00,
        "volume_discount_amt": d.get("volume_discount_amt", 0.00),
        "pro_discount_amt": d.get("pro_discount_amt", 0.00),
        "num_prints": d["num_prints"],
        "regular_price": d["total_cost"] or 0.00,
        "pro_price": d["pro_total_cost"] or 0.00,
        "frame_cost": d["frame_cost"] or 0.00,
        "stretching_fee": d["stretching_fee"] or 0.00,
        "bracer_cost": d["bracer_cost"] or 0.00,
        "upcharge": d["upcharge"] or 0.00,
        "color": d["color"],
        "artist_first": d.get("artist_first", "").strip(),
        "artist_last": d.get("artist_last", "").strip(),
        "title": d.get("title", "").strip(),
        "linked_title": d.get("linked_title", current_title),
        "custom_items": d.get("custom_items", []),
    }


    if d["capture_price"]:
        if d.get("original_height") is None or d.get("original_width") is None:
            messagebox.showerror("Missing Dimensions", "Original height and width are required for Capture.")
            return
        max_dim = max(d["original_height"], d["original_width"])
        draft_items.append({
            "print_type": f"📸 {d['capture_size']}" + (f" ({size_formatted})" if size_formatted else ""),
            "size": "",
            "num_prints": 1,
            "regular_price": d["capture_price"],
            "pro_price": d["capture_price"],
            "canvas_cost": 0.00,
            "pro_canvas_cost": 0.00,
            "frame_cost": 0.00,
            "stretching_fee": 0.00,
            "bracer_cost": 0.00,
            "upcharge": 0.00,
            "color": "#FFD700",
            "linked_title": d.get("linked_title", current_title)
        })
        ###print(f"🧩 Add-on added: {d['print_type']}, linked_title: {d.get('linked_title')}")


    if d["specialty_capture"]:
        if d.get("original_height") is None or d.get("original_width") is None:
            messagebox.showerror("Missing Dimensions", "Original height and width are required for Specialty Capture.")
            return
        draft_items.append({
            "print_type": "✨ Specialty Capture",
            "size": "",
            "num_prints": 1,
            "regular_price": 30.00,
            "pro_price": 30.00,
            "canvas_cost": 0.00,
            "pro_canvas_cost": 0.00,
            "frame_cost": 0.00,
            "stretching_fee": 0.00,
            "bracer_cost": 0.00,
            "upcharge": 0.00,
            "color": "#FFD700",
            "linked_title": d.get("linked_title", current_title)
        })
        ##print(f"🧩 Add-on added: {d['print_type']}, linked_title: {d.get('linked_title')}")


    if d["color_match_var"] and d["color_match_var"].get():
        if d.get("original_height") is None or d.get("original_width") is None:
            messagebox.showerror("Missing Dimensions", "Original height and width are required for Color Match.")
            return
        max_dim = max(d["original_height"], d["original_width"])
        if max_dim < 48:
            label, price = "🎨 Basic Color Match", 80.00
        elif max_dim < 72:
            label, price = "🎨 Basic Color Match – 48\"+", 95.00
        else:
            label, price = "🎨 Basic Color Match – 72\"+", 120.00

        draft_items.append({
            "print_type": label,
            "size": "",
            "num_prints": 1,
            "regular_price": price,
            "pro_price": price,
            "canvas_cost": 0.00,
            "pro_canvas_cost": 0.00,
            "frame_cost": 0.00,
            "stretching_fee": 0.00,
            "bracer_cost": 0.00,
            "upcharge": 0.00,
            "color": "#FF69B4",
            "linked_title": d.get("linked_title", current_title)
        })
        ##print(f"🧩 Add-on added: {d['print_type']}, linked_title: {d.get('linked_title')}")


    if d["monitor_match_var"] and d["monitor_match_var"].get():
        draft_items.append({
            "print_type": "🖥️ Monitor Match",
            "size": "",
            "num_prints": 1,
            "regular_price": 50.00,
            "pro_price": 50.00,
            "canvas_cost": 0.00,
            "pro_canvas_cost": 0.00,
            "frame_cost": 0.00,
            "stretching_fee": 0.00,
            "bracer_cost": 0.00,
            "upcharge": 0.00,
            "color": "#ADD8E6",
            "linked_title": d.get("linked_title", current_title)
        })
        ##print(f"🧩 Add-on added: {d['print_type']}, linked_title: {d.get('linked_title')}")


    if d["complex_wrap_var"] and d["complex_wrap_var"].get():
        draft_items.append({
            "print_type": "🐩 Complex Image Wrap",
            "size": "",
            "num_prints": 1,
            "regular_price": 15.00,
            "pro_price": 15.00,
            "canvas_cost": 0.00,
            "pro_canvas_cost": 0.00,
            "frame_cost": 0.00,
            "stretching_fee": 0.00,
            "bracer_cost": 0.00,
            "upcharge": 0.00,
            "color": "#FFB6C1",  # light pink
            "linked_title": d.get("linked_title", current_title)
        })
        ##print(f"🧩 Add-on added: {d['print_type']}, linked_title: {d.get('linked_title')}")


    
    if d["additional_rounds_var"] and d["additional_rounds_var"].get() > 0:
        rounds = d["additional_rounds_var"].get()
        draft_items.append({
            "print_type": "💻 Additional Color Match Rounds",
            "size": "",
            "num_prints": rounds,
            "regular_price": 30.00,
            "pro_price": 30.00,
            "canvas_cost": 0.00,
            "pro_canvas_cost": 0.00,
            "frame_cost": 0.00,
            "stretching_fee": 0.00,
            "bracer_cost": 0.00,
            "upcharge": 0.00,
            "color": "#FF69B4",
            "linked_title": d.get("linked_title", current_title)
        })
        ##print(f"🧩 Add-on added: {d['print_type']}, linked_title: {d.get('linked_title')}")

    if d["flashdrive_var"] and d["flashdrive_var"].get() > 0:
        count = d["flashdrive_var"].get()
        draft_items.append({
            "print_type": "💿 Flashdrive",
            "size": "",
            "num_prints": count,
            "regular_price": 10.00,
            "pro_price": 10.00,
            "canvas_cost": 0.00,
            "pro_canvas_cost": 0.00,
            "frame_cost": 0.00,
            "stretching_fee": 0.00,
            "bracer_cost": 0.00,
            "upcharge": 0.00,
            "color": "#D3D3D3",  # light gray
            "linked_title": d.get("linked_title", current_title)
        })

    if d["computer_time_var"] and d["computer_time_var"].get() > 0:
        hours = d["computer_time_var"].get()
        rounded_hours = math.ceil(hours * 4) / 4.0
        draft_items.append({
            "print_type": "🕖 Computer Time",
            "size": "",
            "num_prints": rounded_hours,
            "regular_price": 100.00,
            "pro_price": 100.00,
            "canvas_cost": 0.00,
            "pro_canvas_cost": 0.00,
            "frame_cost": 0.00,
            "stretching_fee": 0.00,
            "bracer_cost": 0.00,
            "upcharge": 0.00,
            "color": "#B0C4DE",  # light steel blue
            "linked_title": d.get("linked_title", current_title)
        })



    for var in [d["capture_var"], d["specialty_capture_var"], d["color_match_var"], d["monitor_match_var"], d["complex_wrap_var"]]:
        if var: var.set(False)
    
    
    if d["additional_rounds_var"]:
        d["additional_rounds_var"].set(0)
    if d["flashdrive_var"]:
        d["flashdrive_var"].set(0)
    if d["computer_time_var"]:
        d["computer_time_var"].set(0.0)


    if d["print_type"] != "Add-On Only Order" and (d["canvas_cost"] or d["pro_canvas_cost"]):
        draft_items.append(item)
    update_draft_display()

    if d.get("from_results") and d.get("calculate_results"):
        d["calculate_results"]()

    # 🔁 Restore scroll position (if it existed)
    if previous_scroll:
        results_box.after_idle(lambda: results_box.yview_moveto(previous_scroll[0]))




# In[12]:


def update_draft_display():
    """
    Refreshes the draft invoice display with an itemized breakdown.
    """

    global title_start_indices
    global draft_box, draft_frame
    global current_artist, current_title
    global editable_title_ranges
    editable_title_ranges = {}
    global title_click_regions
    title_click_regions = {}
    global custom_items_by_title



    draft_box.tag_remove("current_focus", "1.0", "end")
    title_start_indices = {}
    draft_box.config(state="normal")
    # Save scroll position and anchor
    try:
        current_scroll = draft_box.yview()
    except Exception:
        current_scroll = (0.0, 1.0)
    draft_box.delete("1.0", tk.END)  # Clear previous entries

    # Define visual styles
    draft_box.tag_configure("artist", font=("Monaco", 18, "bold"), foreground="#00ffcc", justify="center")
    draft_box.tag_configure("title", font=("Monaco", 15, "italic"), foreground="#00ffcc", justify="left")
    draft_box.tag_configure("header", font=("Monaco", 15, "bold"), foreground="#00ffcc", justify="left")
    
    if current_artist:
        draft_box.insert(tk.END, f"{current_artist}\n", "artist")

    for widget in draft_frame.winfo_children():
        if isinstance(widget, tk.Button) and "✖️" in widget.cget("text"):
            widget.destroy()

    #print(f"🖼️ Updating draft display with draft_titles: {draft_titles}")
    unique_index = 0


    prev_title = None
    prev_block_start = None
    
    for title in draft_titles:
        block_start = draft_box.index(tk.END)

        # If we just finished a previous block, store its region now
        if prev_title is not None:
            block_end = block_start
            title_click_regions[prev_title] = {
                "click": (prev_block_start, block_end)
            }
        
        # Update previous title tracker
        prev_title = title
        prev_block_start = block_start

        #print(f"🧾 Drawing section for title: {title}")
        toggle_text = "▼" if title not in collapsed_titles else "▶"
        toggle_btn = make_toggle_button(
            draft_box,
            toggle_text,
            title,
            lambda t=title: toggle_title_visibility(t),
            tag="toggle"
        )
        
        remove_button = make_toggle_button(
            draft_box,
            "✖️",
            title,
            lambda t=title: remove_title_block(t),
            tag="remove"
        )
        
        send_button = make_toggle_button(
            draft_box,
            "✔️",
            title,
            lambda t=title: send_title_block_to_invoice(t),
            tag="send"
        )
        
        # Insert clickable divider and store index for focus range
        divider_start = draft_box.index(tk.END)
        draft_box.insert(tk.END, "     " + ("=" * 40) + "\n", "header")
        divider_line = draft_box.index(f"{divider_start} linestart")
        title_start_indices[title] = (divider_line, f"{divider_line} lineend")
        block_end = draft_box.index(tk.END)


        draft_box.insert(tk.END, "  ")
        draft_box.window_create(tk.END, window=toggle_btn)
        draft_box.insert(tk.END, "  ")
        
        title_entry_widget = create_editable_title_entry(draft_box, title)
        draft_box.window_create(tk.END, window=title_entry_widget)

        draft_box.insert("end", "\n")

        draft_box.insert(tk.END, "  ")
        draft_box.window_create(tk.END, window=remove_button)
        
        draft_box.insert(tk.END, "  ")
        draft_box.window_create(tk.END, window=send_button)

        draft_box.insert("end", "\n")
        
        # Debug: Show that the function is being called
        ##print(f"DEBUG: Updating draft display. Current items: {len(draft_items)}")
    
        if not draft_items:
            draft_box.insert(tk.END, "Draft is empty.\n")
    
        draft_box.update_idletasks()
        
                # Set end of block after all content for this title
        block_end = draft_box.index(tk.END)
        
        # Always store click region for the title block
        if title not in collapsed_titles:
            
            for global_index, item in enumerate(draft_items):
                if item.get("linked_title") != title:
                    continue
                color = item["color"]
                tag_name = f"color_{unique_index}"
                draft_box.tag_configure(tag_name, foreground=color)
    
                title_line = f"(Quantity: {item['num_prints']}) {item['print_type']}"
                if item['size']:
                    title_line += f" - {item['size']}"
                title_line += "\n"
                draft_box.insert(tk.END, title_line, tag_name)

                if item.get("source") == "custom":
                    draft_box.insert(tk.END, f"  Price: ${item['regular_price']:.2f}\n")
                elif any(key in item["print_type"] for key in ["📸", "✨", "🎨", "💻", "🖥️", "🐩", "💿", "🕖"]):
                    draft_box.insert(tk.END, f"  Price: ${item['regular_price']:.2f}\n")
                else:        
                    if item.get("canvas_cost", 0) > 0:
                        draft_box.insert(tk.END, f"  Print Price: ${item['canvas_cost']:.2f}\n")
                    if item.get("pro_canvas_cost", 0):
                        draft_box.insert(tk.END, f"  Professional Print Price: ${item['pro_canvas_cost']:.2f}\n")
                    if item.get("frame_cost", 0):
                        draft_box.insert(tk.END, f"  Frame Wood: ${item['frame_cost']:.2f}\n")
                    if item.get("stretching_fee", 0):
                        draft_box.insert(tk.END, f"  Stretch: ${item['stretching_fee']:.2f}\n")
                    if item.get("bracer_cost", 0):
                        draft_box.insert(tk.END, f"  Bracer Wood: ${item['bracer_cost']:.2f}\n")
                    if item.get("upcharge", 0):
                        draft_box.insert(tk.END, f"  ≥ 72\" Upcharge: ${item['upcharge']:.2f}\n")
                    #if item["regular_price"] > 0:
                        #draft_box.insert(tk.END, f"  Regular Total: ${item['regular_price']:.2f}\n")
                    #if item["pro_price"] > 0:
                        #draft_box.insert(tk.END, f"  Professional Total: ${item['pro_price']:.2f}\n")
        
            # 🟩 Create button row BEFORE divider
                draft_box.config(state="normal")
                draft_box.update_idletasks()

                def handle_remove(idx=global_index):
                    item = draft_items[idx]
                
                    # If it's a custom draft line, remove from custom_items_by_title too
                    if item.get("source") == "custom":
                        t = item.get("linked_title")
                        cid = item.get("custom_id")
                
                        if t in custom_items_by_title and cid:
                            custom_items_by_title[t] = [
                                ci for ci in custom_items_by_title[t]
                                if ci.get("id") != cid
                            ]
                
                    draft_items.pop(idx)
                    update_draft_display()

                    
                remove_button = tk.Button(
                    draft_box,
                    text="✖️",
                    font=("Avenir Next", 10),
                    command=handle_remove,
                    bg="#222222", fg="white", activebackground="#333333",
                    bd=0, highlightthickness=0, padx=2, pady=0
                )
                draft_box.insert(tk.END, "  ")  
                draft_box.window_create(tk.END, window=remove_button)
            
                draft_box.insert("end", "\n")
        
                # Divider comes *after* the button row
                draft_box.insert(tk.END, "=" * 40 + "\n", tag_name)
                unique_index += 1
    
        # Capture the final title block
        if prev_title is not None:
            block_end = draft_box.index(tk.END)
            title_click_regions[prev_title] = {
                "click": (prev_block_start, block_end)
            }

    
            
    draft_box.config(state="disabled")
    draft_box.update_idletasks()

    # Restore scroll unless we've changed to a new title block
    if current_title in title_start_indices:
        start, end = title_start_indices[current_title]
        start_line = int(start.split('.')[0])
    
        # Only reset view if current title would scroll out of range
        current_top_line = int(draft_box.index("@0,0").split('.')[0])
        current_bottom_line = int(draft_box.index("@0,%d" % draft_box.winfo_height()).split('.')[0])
    
        if not (start_line >= current_top_line and start_line <= current_bottom_line):
            draft_box.see(start)  # Bring into view
    else:
        # If title was removed or no longer in view, fallback
        draft_box.yview_moveto(current_scroll[0])

    


# In[13]:


def create_editable_title_entry(parent, title):
    
    title_var = tk.StringVar(value=title)
    entry = tk.Entry(
        parent,
        textvariable=title_var,
        font=("Monaco", 15, "italic"),
        fg="#00ffcc",
        bg="#222222",
        bd=0,
        insertbackground="#00ffcc",
        highlightthickness=0,
        width=len(title) + 2
    )

    def update_title(event, original_title=title):
        global draft_titles, draft_items, current_title, title_entry

        new_title = entry.get().strip()
        #print(f"🔁 Attempting to rename title: '{original_title}' → '{new_title}'")

        if not new_title or new_title == original_title:
            #print"⛔ No change or blank title.")
            return

        if new_title in draft_titles:
            #print(f"❗ Conflict detected: '{new_title}' already in draft_titles")
        
            # Reset entry to original title
            entry.delete(0, tk.END)
            entry.insert(0, original_title)
            return

        #print(f"✏️ Renaming title: '{original_title}' → '{new_title}' (no conflict)")

        try:
            idx = draft_titles.index(original_title)
            draft_titles[idx] = new_title
        except ValueError:
            #print"⚠️ Original title not found — appending new title.")
            draft_titles.append(new_title)

        draft_titles = list(dict.fromkeys(draft_titles))

        for item in draft_items:
            if item.get("linked_title") == original_title:
                item["linked_title"] = new_title

        current_title = new_title
        title_entry.delete(0, tk.END)
        title_entry.insert(0, new_title)
        update_draft_display()

        def animate_success_flash():
            if new_title in title_start_indices:
                start, end = title_start_indices[new_title]
                #print(f"🌿 Success flash for: {new_title} from {start} to {end}")
                draft_box.tag_add("success_flash", start, end)
                draft_box.tag_configure("success_flash", background="#32CD32")
                draft_box.update_idletasks()
                draft_box.after(25, lambda: draft_box.tag_remove("success_flash", start, end))

        draft_box.after(100, animate_success_flash)

    entry.bind("<Return>", update_title)
    entry.bind("<FocusOut>", update_title)

    return entry


# In[14]:


def edit_artist():
    artist_first_entry.focus_set()  # Optional: could use .icursor(tk.END) to place cursor

def remove_artist():
    global current_artist
    current_artist = ""
    artist_first_entry.delete(0, tk.END)
    artist_last_entry.delete(0, tk.END)
    update_draft_display()

def edit_title():
    title_entry.focus_set()

def remove_title_block(title):
    global draft_items, draft_titles
    draft_items = [item for item in draft_items if item.get("linked_title") != title]
    if title in draft_titles:
        draft_titles.remove(title)
    update_draft_display()

def remove_title_from_invoice(title):
    global invoice_items
    invoice_items = [item for item in invoice_items if item.get("linked_title") != title]
    update_invoice_display(apply_tax_var.get())

def send_current_title_to_invoice():
    global draft_items, invoice_items, current_title

    if not current_title:
        return  # No title to match against

    # Move all draft items with matching linked_title
    matched_items = [item for item in draft_items if item.get("linked_title") == current_title]
    draft_items[:] = [item for item in draft_items if item.get("linked_title") != current_title]
    
    invoice_items.extend(matched_items)
    
    update_invoice_display(apply_tax=apply_tax_var.get())
    update_draft_display()
    #print(f"🛫 Sending all items linked to title: {current_title}")

def send_title_block_to_invoice(title):
    global draft_items, invoice_items, draft_titles
    #print(f"📤 Sending all items for title: '{title}' to invoice")
    #print(f"  Before send — draft_items count: {len(draft_items)}")
    matched = [item for item in draft_items if item.get("linked_title") == title]
    draft_items = [item for item in draft_items if item.get("linked_title") != title]
    invoice_items.extend(matched)
    #print(f"  After send — invoice_items count: {len(invoice_items)}")
    if title in draft_titles:
        draft_titles.remove(title)
    update_invoice_display(apply_tax=apply_tax_var.get())
    update_draft_display()

def toggle_title_visibility(title):
    if title in collapsed_titles:
        collapsed_titles.remove(title)
    else:
        collapsed_titles.add(title)

    update_draft_display()
    update_invoice_display(apply_tax_var.get())


def handle_click_in_draft_box(event, calculate_results_fn):
    global current_title
    clicked_index = draft_box.index(f"@{event.x},{event.y}")
    #print(f"🖱️ Clicked at {clicked_index}")

    for title, region in title_click_regions.items():
        start, end = region["click"]
        if draft_box.compare(clicked_index, ">=", start) and draft_box.compare(clicked_index, "<", end):
            if current_title != title:
                current_title = title
                title_entry.delete(0, tk.END)
                title_entry.insert(0, title)
                #print(f"📍 Focus switched to title: {title}")

                def highlight_after_redraw():
                    draft_box.tag_remove("current_focus", "1.0", "end")
                    draft_box.tag_configure("current_focus", background="#333355")
                    draft_box.tag_add("current_focus", start, end)

                calculate_results_fn()
                draft_box.after(100, highlight_after_redraw)
            break


button_states = {}

def make_toggle_button(master, label, title, callback, tag=""):
    """
    Creates a styled button that shows SMS orange while pressed.
    - Normal state: black icon on dark gray
    - Pressed: orange icon on darker gray
    - Reverts immediately after release
    """
    return tk.Button(
        master,
        text=label,
        font=("Avenir Next", 10),
        command=callback,
        bg="#222222",             # Normal background
        fg="black",               # Normal text color
        activebackground="#373737",  # Background while pressed
        activeforeground="#C93214",  # Orange while pressed
        bd=0,
        highlightthickness=0,
        padx=4,
        pady=0
    )



# In[15]:


def clear_draft():
    """
    Clears all items from the draft list.
    """
    global draft_items, draft_titles, current_artist, current_title
    draft_items.clear()
    draft_titles.clear()
    #current_artist = ""
    current_title = ""
    update_draft_display()


# In[16]:


def on_artist_or_title_change(event=None):
    if invoice_items:
        update_invoice_display(apply_tax=apply_tax_var.get())


# In[17]:


def send_to_invoice(index=None):
    """
    Moves draft items to the invoice box.
    - If index is None: move all items.
    - If index is a number: move only that one item.
    """
    global draft_items, invoice_items, current_artist, current_title
    
    apply_tax = apply_tax_var.get()

    if index is None:
        invoice_items.extend(draft_items)
        update_invoice_display(apply_tax=apply_tax)
        draft_items.clear()
        update_draft_display()
        #current_artist = ""
        current_title = ""
    else:
        if 0 <= index < len(draft_items):
            invoice_items.append(draft_items.pop(index))
        update_invoice_display(apply_tax=apply_tax)
        update_draft_display()
        


# In[18]:


def update_invoice_display(apply_tax=None):
    global invoice_box, invoice_frame, apply_tax_var, apply_card
    global current_artist, current_title

    # 🔥 ALWAYS sync current_artist from the GUI textboxes:
    try:
        full_name = (artist_first_entry.get().strip() + " " + artist_last_entry.get().strip()).strip()
        if full_name:
            current_artist = full_name
    except NameError:
        # Entries may not be created yet during initial setup
        pass

    if apply_tax is None:
        apply_tax = apply_tax_var.get()

    invoice_box.config(state="normal")
    try:
        current_scroll = invoice_box.yview()
    except Exception:
        current_scroll = (0.0, 1.0)
    invoice_box.delete("1.0", tk.END)
        
    tax_rate = 0.07 if apply_tax else 0.00
    final_subtotal = 0.0
    final_tax = 0.0

    invoice_box.tag_configure("artist", font=("Monaco", 18, "bold"), foreground="#00ffcc", justify="center")
    invoice_box.tag_configure("title", font=("Monaco", 15, "italic"), foreground="#00ffcc", justify="left")
    invoice_box.tag_configure("header", font=("Monaco", 15, "bold"), foreground="#00ffcc", justify="left")

    if current_artist:
        invoice_box.insert(tk.END, f"{current_artist}\n", "artist")

    if not invoice_items:
        invoice_box.insert(tk.END, "Invoice is empty.\n")
        invoice_box.config(state="disabled")
        return

    show_regular = price_var_regular.get()
    show_pro = price_var_pro.get()

    items_by_title = {}
    for item in invoice_items:
        title = item.get("linked_title") or "Untitled"
        items_by_title.setdefault(title, []).append(item)

    unique_index = 0
    for title_index, (title, items) in enumerate(items_by_title.items()):
        invoice_box.insert(tk.END, "     " + ("=" * 40) + "\n", "header")
    
        invoice_box.insert(tk.END, "  ")
        collapsed = title in collapsed_titles
        toggle_symbol = "▼" if not collapsed else "▶"
        toggle_btn = make_toggle_button(invoice_box, toggle_symbol, title, lambda t=title: toggle_title_visibility(t), tag="toggle")
        invoice_box.window_create(tk.END, window=toggle_btn)
        invoice_box.insert(tk.END, "  ")
        invoice_box.insert(tk.END, f"{title}\n", "title")
    
        remove_title_btn = tk.Button(invoice_box, text="✖️", font=("Avenir Next", 10),
                                     command=lambda t=title: remove_title_from_invoice(t),
                                     bg="#222222", fg="white", activebackground="#333333",
                                     bd=0, highlightthickness=0, padx=2, pady=0)
        invoice_box.insert(tk.END, "  ")
        invoice_box.window_create(tk.END, window=remove_title_btn)
        invoice_box.insert(tk.END, "\n")
    
        for item_index, item in enumerate(items):
            # Always compute totals
            unit_price = round(item["regular_price"], 2)
            quantity = item["num_prints"]
            subtotal = round(unit_price * quantity, 2)
            final_subtotal = round(final_subtotal + subtotal, 2)

    
            if collapsed:
                continue  # Skip rendering item details
    
            # Rendering (only if not collapsed)
            index = title_index * 100 + item_index
            tag_name = f"color_{index}"
            invoice_box.tag_configure(tag_name, foreground=item["color"])
    
            is_print = item.get("canvas_cost", 0) > 0 or item.get("pro_canvas_cost", 0) > 0

    
            line = f"(Quantity: {quantity}) {item['print_type']}"
            if item["size"]:
                line += f" - {item['size']}"
            line += "\n"
            invoice_box.insert(tk.END, line, tag_name)
    
            if is_print:
                if show_regular and item.get("canvas_cost", 0) > 0:
                    invoice_box.insert(tk.END, f"  Regular Print Price: ${item['canvas_cost']:.2f}\n")
                if show_pro and item.get("pro_canvas_cost", 0) > 0:
                    invoice_box.insert(tk.END, f"  Professional Print Price: ${item['pro_canvas_cost']:.2f}\n")
                if item.get("frame_cost", 0) > 0:
                    invoice_box.insert(tk.END, f"  Frame Wood: ${item['frame_cost']:.2f}\n")
                if item.get("stretching_fee", 0) > 0:
                    invoice_box.insert(tk.END, f"  Stretch: ${item['stretching_fee']:.2f}\n")
                if item.get("bracer_cost", 0) > 0:
                    invoice_box.insert(tk.END, f"  Bracer Wood: ${item['bracer_cost']:.2f}\n")
                if item.get("upcharge", 0) > 0:
                    invoice_box.insert(tk.END, f"  ≥ 72\" Upcharge: ${item['upcharge']:.2f}\n")
                invoice_box.insert(tk.END, "     " + ("⋯⋯⋯⋯⋯" * 3) + "\n")
                if show_regular and item.get("regular_price", 0) > 0:
                    invoice_box.insert(tk.END, f"  Regular Total: ${item['regular_price']:.2f}\n")
                if show_pro and item.get("pro_price", 0) > 0:
                    invoice_box.insert(tk.END, f"  Professional Total: ${item['pro_price']:.2f}\n")
            elif quantity == 1:
                invoice_box.insert(tk.END, f"  Price: ${unit_price:.2f}\n")
            else:
                invoice_box.insert(tk.END, f"  Unit Price: ${unit_price:.2f} each\n")
                invoice_box.insert(tk.END, f"  Subtotal: ${subtotal:.2f}\n")
    
            invoice_box.insert(tk.END, "=" * 40 + "\n", tag_name)
            unique_index += 1


    # --- Step 2: Built-in volume / pro discounts + custom discounts ---

    # Savings from built-in discounts (already stored per item)
    total_volume_discount = sum(item.get("volume_discount_amt", 0.00) for item in invoice_items)
    total_pro_discount = sum(item.get("pro_discount_amt") or 0.00 for item in invoice_items)

    volume_savings = abs(total_volume_discount)
    pro_savings = abs(total_pro_discount)

    # Custom discounts
    dollar_discount = custom_discount_dollar_var.get()
    percent_discount = custom_discount_var.get()

    # Percent discount is applied on the FULL Subtotal
    percent_base = final_subtotal
    percent_discount_amt = percent_base * (percent_discount / 100)

    # Discounted Subtotal subtracts ALL discount lines:
    # volume, professional, flat, and custom percent.
    discounted_subtotal = (
        final_subtotal
        - volume_savings
        - pro_savings
        - dollar_discount
        - percent_discount_amt
    )
    if discounted_subtotal < 0:
        discounted_subtotal = 0.0

    # --- Step 3: Apply card fee + tax (tax includes card fee) ---
    
    # Normalize / protect
    discounted_subtotal = round(discounted_subtotal, 2)
    
    # 1) Card fee is computed BEFORE tax
    card_fee = 0.0
    if apply_card.get():
        card_fee = round(discounted_subtotal * 0.03, 2)
    
    # 2) Tax base includes card fee (so you tax the fee)
    taxable_base = round(discounted_subtotal + card_fee, 2)
    
    final_tax = 0.0
    if tax_rate > 0:
        final_tax = round(taxable_base * tax_rate, 2)
    
    # 3) Final total
    final_total = round(discounted_subtotal + card_fee + final_tax, 2)

    # --- DISPLAY SUMMARY (in this order) ---
    
    invoice_box.insert(tk.END, f"Subtotal: ${final_subtotal:.2f}\n")

    
    
    
    if volume_savings > 0:
        invoice_box.insert(tk.END, f"Volume Discount Savings: -${volume_savings:.2f}\n")
    if pro_savings > 0:
        invoice_box.insert(tk.END, f"Professional Discount Savings: -${pro_savings:.2f}\n")

    #print(f"After invoice_box.insert(tk.END, Volume Discount Savings...): {volume_savings=}")
    #print(f"After invoice_box.insert(tk.END, Professional Discount Savings..: {pro_savings=}")

    
    if dollar_discount > 0:
        invoice_box.insert(tk.END, f"Flat Discount: -${dollar_discount:.2f}\n")
    
    if percent_discount > 0:
        invoice_box.insert(tk.END, f"Custom Discount ({percent_discount:.2f}%): -${percent_discount_amt:.2f}\n")
    
    invoice_box.insert(tk.END, f"Discounted Subtotal: ${discounted_subtotal:.2f}\n")

    if apply_card.get():
        invoice_box.insert(tk.END, f"Card Fee (3%): ${card_fee:.2f}\n")
    
    if tax_rate > 0:
        invoice_box.insert(tk.END, f"Sales Tax ({int(tax_rate * 100)}%): ${final_tax:.2f}\n")

    invoice_box.insert(tk.END, f"Total Due: ${final_total:.2f}\n", "header")
    
    invoice_summary = {
        "final_subtotal": final_subtotal,
        "dollar_discount": dollar_discount,
        "percent_discount": percent_discount,
        "percent_discount_amt": percent_discount_amt,
        "discounted_subtotal": discounted_subtotal,
        "final_tax": final_tax,
        "final_card_fee": card_fee,
        "final_total": final_total,
        "volume_savings": volume_savings,
        "pro_savings": pro_savings,
    }

    # This is the order it'll show up in in the PDF and QBO

    invoice_summary["summary_lines"] = [
        ("Subtotal", final_subtotal),
        ("Volume Discount", -volume_savings),
        ("Professional Discount", -pro_savings),
        ("Flat Discount", -dollar_discount),
        (f"Custom Discount ({percent_discount:.2f}%)", -percent_discount_amt),
        ("Discounted Subtotal", discounted_subtotal),
        ("Card Fee (3%)", card_fee),
        (f"Sales Tax ({int(tax_rate * 100)}%)", final_tax),
        ("Total Due", final_total),
    ]

    # Group items by title before generating invoice_data
    items_by_title = {}
    for item in invoice_items:
        title = item.get("linked_title") or "Untitled"
        items_by_title.setdefault(title, []).append(item)
    
    global invoice_prices
    invoice_prices = {
        "artist": current_artist,
        "items_by_title": items_by_title,
        "summary": invoice_summary,
    }

    
    invoice_box.yview_moveto(current_scroll[0])
    invoice_box.config(state="disabled")
    invoice_box.update_idletasks()

    return invoice_prices


# In[19]:


from docx import Document
from docx.shared import Pt, Inches
from docx2pdf import convert
from docx.shared import Pt, RGBColor
import os
from datetime import datetime

def generate_invoice_docx(invoice_data, output_path="Generated_Invoice.docx", apply_tax=True, apply_card_fee=True):
    template_path =  # Change this on new system ⚠️⚠️⚠️
    doc = Document(template_path)

    from docx.shared import RGBColor, Pt
    
    #print("==== SUMMARY LINES ====")
    #print(invoice_prices["summary"]["summary_lines"])
    
    #print("\n==== FINAL TOTAL ====")
    #print(invoice_prices["summary"]["final_total"])

    # --- Helpers ---
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    
    def add_aligned_line(paragraph, left_text, right_text, tab_pos=Inches(6.25), add_dots=True):
        paragraph.paragraph_format.tab_stops.clear_all()
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            tab_pos,
            alignment=WD_TAB_ALIGNMENT.RIGHT,
            leader=WD_TAB_LEADER.DOTS if add_dots else WD_TAB_LEADER.SPACES
        )
        r_left = paragraph.add_run(left_text)
        paragraph.add_run("\t")
        r_right = paragraph.add_run(right_text)
        return r_left, r_right
    
    def add_indented_price_line(label, amount):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)  # Indent from left
        add_aligned_line(p, f"  {label}:", f"${amount:.2f}", tab_pos=Inches(4.5))  # Slightly shorter tab stop
    
    
    def add_price_line(label, amount):
        p = doc.add_paragraph()
        add_aligned_line(p, f"  {label}:", f"${amount:.2f}")

    doc.add_paragraph("")
    doc.add_paragraph("")
    
    # 🎨 Artist name (left) + Date (right)
    p_header = doc.add_paragraph()
    r_artist, r_date = add_aligned_line(
        p_header,
        invoice_data.get("artist", "Unknown Artist"),
        datetime.now().strftime("%Y %B, %d"),
        add_dots=False
    )

    
    # Apply styles
    for run, size in [(r_artist, 28), (r_date, 18)]:
        run.font.name = "Avenir Next"
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0x1C, 0x52, 0x3F)

    
    doc.add_paragraph("")

    # Summary
    summary = invoice_prices.get("summary", {})
    summary_lines = summary.get("summary_lines", [])
    
    items_by_title = invoice_prices.get("items_by_title", {})
    use_pro = invoice_prices.get("use_pro", False)

    first_block = True

    for title, items in items_by_title.items():
        if not first_block:
            doc.add_paragraph("=" * 50)
        first_block = False
        doc.add_paragraph(title, style="Heading 2")

        for item in items:
            unit_price = round(item["pro_price"] if use_pro else item["regular_price"], 2)
            quantity = item["num_prints"]
            
            is_service = (
                any(tag in item["print_type"] for tag in ["📸", "✨", "🎨", "💻", "🖥️", "🐩", "🕖", "💿"])
                or item.get("source") == "custom"
            )

            size_text = f" - {item['size']}" if item["size"] else ""
            item_line = f"{quantity} x {item['print_type']}{size_text}"

            print_type_raw = item.get("print_type", "")
            pt = print_type_raw.lower()
            
            # Determine material label once
            material_label = "Canvas" if "canvas" in pt else "Paper"
            
            # Collapse rules: paper stocks + unstretched canvas => single line only
            is_paper_collapse = ("enhanced matte" in pt) or ("photorag" in pt) or ("watercolor" in pt)
            is_unstretched_collapse = "unstretched canvas" in pt
            collapse_line = is_paper_collapse or is_unstretched_collapse
            
            if is_service:
                p = doc.add_paragraph()
                add_aligned_line(
                    p,
                    item_line.replace("🎨","").replace("📸","").replace("✨","")
                             .replace("💻","").replace("🖥️","").replace("🐩","")
                             .replace("🕖","").replace("💿","").strip(),
                    f"${unit_price * quantity:.2f}"
                )
            else:
                if collapse_line:
                    # One clean line: "<qty> x <type> - <size> ........................ $price"
                    p = doc.add_paragraph()
                    add_aligned_line(p, item_line, f"${unit_price * quantity:.2f}")
                    # No material breakdown or addons for these
                else:
                    # Keep existing breakdown for stretched canvas etc.
                    doc.add_paragraph(item_line)
            
                    if use_pro and item.get("pro_canvas_cost", 0.0) > 0:
                        add_indented_price_line(f"Pro {material_label}", item["pro_canvas_cost"])
                    elif item.get("canvas_cost", 0.0) > 0:
                        add_indented_price_line(material_label, item["canvas_cost"])
            
                    if item.get("frame_cost", 0) > 0:
                        add_indented_price_line("Frame Wood", item["frame_cost"])
                    if item.get("stretching_fee", 0) > 0:
                        add_indented_price_line("Stretch", item["stretching_fee"])
                    if item.get("bracer_cost", 0) > 0:
                        add_indented_price_line("Bracer Wood", item["bracer_cost"])
                    if item.get("upcharge", 0) > 0:
                        add_indented_price_line("≥ 72\" Upcharge", item["upcharge"])
            
                    add_price_line("Print Total", round((unit_price * quantity), 2))
                    doc.add_paragraph("")



    # Summary Divider
    doc.add_paragraph("=" * 70)
    
    # Summary
    summary = invoice_data.get("summary", {})
    subtotal = None
    discounted_subtotal = None
    
    for label, amount in summary.get("summary_lines", []):
        lower_label = label.lower()
        is_discount_line = any(key in lower_label for key in ["discount", "savings"])
    
        # Skip tax line if apply_tax is False
        if "tax" in lower_label and not apply_tax:
            continue
    
        # Skip card fee line if apply_card_fee is False
        if "card fee" in lower_label and not apply_card_fee:
            continue
    
        # Capture subtotals for later comparison
        if "subtotal" in lower_label and "discounted" not in lower_label:
            subtotal = round(amount, 2)
        elif "discounted subtotal" in lower_label:
            discounted_subtotal = round(amount, 2)
    
        # Skip 0-value discount/savings lines
        if is_discount_line and abs(amount) < 0.005:
            continue
    
        # Skip discounted subtotal if it matches the original subtotal
        if "discounted subtotal" in lower_label and subtotal == discounted_subtotal:
            continue

        # 🧠 Skip both subtotal and discounted subtotal if subtotal == final total
        final_total = round(summary.get("final_total", 0.0), 2)
        if ("subtotal" in lower_label) and (subtotal == final_total):
            continue
        if ("total" in lower_label) and (subtotal == final_total):
            continue
    
        # Render the line
        add_price_line(label, amount)

    
    # Total (bold and larger)
    p_total = doc.add_paragraph()
    p_total.paragraph_format.alignment = 2  # Right align
    total_run = p_total.add_run(f"Total Due: ${summary.get('final_total', 0):.2f}")
    total_run.font.size = Pt(20)
    total_run.bold = True
    total_run.font.color.rgb = RGBColor(0x25, 0x52, 0x90)


    doc.save(output_path)
    
    # Convert to PDF and open it
    pdf_path = output_path.replace(".docx", ".pdf")
    try:
        convert(output_path)
        os.system(f"open '{pdf_path}'")  # macOS
    except Exception as e:
        print(f"PDF conversion or opening failed: {e}")




# In[20]:


def remove_invoice_item(index):
    """
    Removes an individual item from the invoice list and updates the display.
    """
    global invoice_items
    if 0 <= index < len(invoice_items):
        del invoice_items[index] 
        update_invoice_display()


# In[21]:


def clear_invoice():
    """
    Clears all items from the invoice list.
    """
    global invoice_items, current_artist, current_title
    invoice_items.clear()
    #current_artist = ""
    current_title = ""
    update_invoice_display()



# In[22]:


from dotenv import load_dotenv
import requests
import json
from subprocess import run

def refresh_access_token():
    refresh_token = os.getenv("REFRESH_TOKEN")
    client_id = os.getenv("CLIENT_ID")
    client_secret = os.getenv("CLIENT_SECRET")


    token_url = "https://oauth.platform.intuit.com/oauth2/v1/tokens/bearer"

    headers = {
        "Accept": "application/json",
        "Content-Type": "application/x-www-form-urlencoded"
    }

    payload = {
        "grant_type": "refresh_token",
        "refresh_token": refresh_token
    }

    try:
        response = requests.post(
            token_url,
            headers=headers,
            auth=(client_id, client_secret),
            data=payload
        )
    except requests.RequestException as e:
        logging.error("TOKEN_REFRESH_REQUEST_EXCEPTION %s", str(e))
        print(f"❌ Token refresh request failed: {e}")
        messagebox.showerror(
            "QuickBooks Connection Error",
            f"Token refresh request failed.\n\n{e}\n\nLog file: {LOG_PATH}"
        )
        return None

    if response.status_code == 200:
        tokens = response.json()
        new_access_token = tokens["access_token"]

        # refresh_token may be omitted sometimes; don't blow away the old one
        new_refresh_token = tokens.get("refresh_token") or refresh_token

        update_env({
            "ACCESS_TOKEN": new_access_token,
            "REFRESH_TOKEN": new_refresh_token
        })

        return new_access_token

    else:
        # Log the failure with intuit_tid for support
        log_qbo_error("token_refresh", response)

        tid = get_intuit_tid(response)
        messagebox.showerror(
            "QuickBooks Token Validation Failed",
            "QuickBooks authorization failed.\n\n"
            f"Intuit TID: {tid}\n"
            f"Log file: {LOG_PATH}\n\n"
            "If the problem persists, contact support:\n"
            "828-318-2202\nbenjaminzeidell@gmail.com"
        )
        return None


def update_env(new_vars):
    lines = []
    with open(".env", "r") as f:
        lines = f.readlines()

    for i, line in enumerate(lines):
        for key, val in new_vars.items():
            if line.startswith(key + "="):
                lines[i] = f"{key}={val}\n"

    with open(".env", "w") as f:
        f.writelines(lines)


def qbo_escape_query_string(s: str) -> str:
    # Escape backslashes first, then single quotes for QBO query strings
    return (s or "").replace("\\", "\\\\").replace("'", "\\'")



def create_customer(first_name, last_name, access_token):
    raw_display_name = f"{first_name} {last_name}".strip()
    clean_display_name = strip_emoji(raw_display_name).strip()

    if not clean_display_name:
        messagebox.showerror(
            "QuickBooks Error",
            "Artist first and/or last name is required."
        )
        return None

    escaped_for_query = qbo_escape_query_string(clean_display_name)
    query = f"SELECT * FROM Customer WHERE DisplayName = '{escaped_for_query}'"


    # --- Lookup ---
    try:
        response = requests.get(
            f"{BASE_URL}/v3/company/{REALM_ID}/query",
            headers={
                "Authorization": f"Bearer {access_token}",
                "Accept": "application/json",
            },
            params={"query": query},
        )
    except requests.RequestException as e:
        logging.error("CUSTOMER_LOOKUP_REQUEST_EXCEPTION %s", str(e))
        messagebox.showerror(
            "QuickBooks Connection Error",
            "Could not contact QuickBooks to look up the customer.\n\n"
            f"Error: {e}"
            "If the problem persists, contact support:\n"
            "828-318-2202\nbenjaminzeidell@gmail.com"
        )
        return None

    if response.status_code == 200:
        customers = response.json().get("QueryResponse", {}).get("Customer", []) or []
        if customers:
            return customers[0].get("Id")
    else:
        log_qbo_error("customer_lookup", response, extra={"display_name": clean_display_name})
        # don't messagebox here; we'll try creation next


    # --- Create ---
    customer_data = {
        "GivenName": first_name or clean_display_name,
        "FamilyName": last_name or "",
        "DisplayName": clean_display_name,
    }

    try:
        response = requests.post(
            f"{BASE_URL}/v3/company/{REALM_ID}/customer",
            headers={
                "Authorization": f"Bearer {access_token}",
                "Content-Type": "application/json",
                "Accept": "application/json",
            },
            json=customer_data,
        )
    except requests.RequestException as e:
        logging.error("CUSTOMER_CREATE_REQUEST_EXCEPTION %s", str(e))
        messagebox.showerror(
            "QuickBooks Connection Error",
            "Could not contact QuickBooks to create the customer.\n\n"
            f"Error: {e}"
            "If the problem persists, contact support:\n"
            "828-318-2202\nbenjaminzeidell@gmail.com"
        )
        return None

    if response.status_code == 200:
        return response.json().get("Customer", {}).get("Id")
    else:
        log_qbo_error("customer_create", response, extra={"display_name": clean_display_name})
        tid = get_intuit_tid(response)
        messagebox.showerror(
            "QuickBooks Error",
            "QuickBooks rejected the customer creation request.\n\n"
            f"Intuit TID: {tid}\n"
            f"Log file: {LOG_PATH}"
            "If the problem persists, contact support:\n"
            "828-318-2202\nbenjaminzeidell@gmail.com"
        )
        return None




def get_or_create_item(item_name, access_token):
    if not item_name or not str(item_name).strip():
        messagebox.showerror(
            "QuickBooks Error",
            "Invalid item name. Cannot create invoice line."
        )
        return None

    item_name = item_name.strip()
    
    escaped_name = qbo_escape_query_string(item_name)
    query = f"SELECT * FROM Item WHERE Name = '{escaped_name}'"


    # --- Lookup ---
    try:
        response = requests.get(
            f"{BASE_URL}/v3/company/{REALM_ID}/query",
            headers={
                "Authorization": f"Bearer {access_token}",
                "Accept": "application/json",
            },
            params={"query": query},
        )
    except requests.RequestException as e:
        logging.error("ITEM_LOOKUP_REQUEST_EXCEPTION %s", str(e))
        messagebox.showerror(
            "QuickBooks Connection Error",
            "Could not contact QuickBooks to look up an item.\n\n"
            f"Error: {e}"
            "If the problem persists, contact support:\n"
            "828-318-2202\nbenjaminzeidell@gmail.com"
        )
        return None

    if response.status_code == 200:
        items = response.json().get("QueryResponse", {}).get("Item", []) or []
        if items:
            return {"value": items[0]["Id"]}
    else:
        log_qbo_error("item_lookup", response, extra={"item_name": item_name})

    # --- Create ---
    item_data = {
        "Name": item_name,
        "Type": "Service",
        "IncomeAccountRef": {
            "name": "Sales of Product Income",
            "value": "79"
        }
    }

    try:
        response = requests.post(
            f"{BASE_URL}/v3/company/{REALM_ID}/item",
            headers={
                "Authorization": f"Bearer {access_token}",
                "Content-Type": "application/json",
                "Accept": "application/json",
            },
            json=item_data,
        )
    except requests.RequestException as e:
        logging.error("ITEM_CREATE_REQUEST_EXCEPTION %s", str(e))
        messagebox.showerror(
            "QuickBooks Connection Error",
            "Could not contact QuickBooks to create an item.\n\n"
            f"Error: {e}"
            "If the problem persists, contact support:\n"
            "828-318-2202\nbenjaminzeidell@gmail.com"
        )
        return None

    if response.status_code == 200:
        item_id = response.json().get("Item", {}).get("Id")
        return {"value": item_id} if item_id else None
    else:
        log_qbo_error("item_create", response, extra={"item_name": item_name})
        tid = get_intuit_tid(response)
        messagebox.showerror(
            "QuickBooks Error",
            f"QuickBooks rejected the item '{item_name}'.\n\n"
            f"Intuit TID: {tid}\n"
            f"Log file: {LOG_PATH}"
            "If the problem persists, contact support:\n"
            "828-318-2202\nbenjaminzeidell@gmail.com"
        )
        return None


import re

def strip_emoji(text):
    emoji_pattern = re.compile(
        "["
        "\U0001F600-\U0001F64F"  # Emoticons
        "\U0001F300-\U0001F5FF"  # Symbols & Pictographs
        "\U0001F680-\U0001F6FF"  # Transport & Map
        "\U0001F1E0-\U0001F1FF"  # Flags
        "\U00002500-\U00002BEF"  # Box Drawing + More
        "\U00002702-\U000027B0"
        "\U000024C2-\U0001F251"
        "]+", flags=re.UNICODE
    )
    return emoji_pattern.sub(r'', text)


# In[23]:


def send_to_quickbooks():

    # --- 1. Refresh token and update global ---


    new_token = refresh_access_token()
    if not new_token:
        return

    global ACCESS_TOKEN, invoice_items, current_artist, current_title, apply_card
    ACCESS_TOKEN = new_token

    # Nothing to send? bail out cleanly
    if not invoice_items:
        # print("❌ No items to send.")
        return

    # --- 2. Customer info from last invoice item ---

    first = invoice_items[-1].get("artist_first", "")
    last = invoice_items[-1].get("artist_last", "")
    
    if not first and not last:
        messagebox.showerror(
            "QuickBooks Error",
            "Artist first and/or last name is required before sending to QuickBooks."
        )
        return
    
    customer_id = create_customer(first, last, ACCESS_TOKEN)


    if not customer_id:
        messagebox.showerror(
            "QuickBooks Error",
            "Could not create or find a customer in QuickBooks. "
            "Please check the artist name and try again."
        )
        return

    # --- 3. Refresh invoice summary (used by PDF + QB) ---

    update_invoice_display()  # sets global invoice_prices
    
    tax_code_ref = {"value": "TAX"} if apply_tax_var.get() else {"value": "NON"}

    summary = invoice_prices.get("summary", {}) or {}

    volume_savings       = summary.get("volume_savings", 0.0) or 0.0
    pro_savings          = summary.get("pro_savings", 0.0) or 0.0
    flat_discount        = summary.get("dollar_discount", 0.0) or 0.0
    percent_discount_amt = summary.get("percent_discount_amt", 0.0) or 0.0
    discounted_subtotal  = summary.get("discounted_subtotal", 0.0) or 0.0
    tax_amount           = summary.get("final_tax", 0.0) or 0.0
    card_fee             = summary.get("final_card_fee", 0.0) or 0.0
    final_total          = summary.get("final_total", 0.0) or 0.0

    # --- 4. Generate the PDF as before ---

    items_by_title = {}
    for item in invoice_items:
        title = item.get("linked_title") or "Untitled"
        items_by_title.setdefault(title, []).append(item)

    invoice_docx_path = "Generated_Invoice.docx"
    generate_invoice_docx(
        invoice_prices,
        output_path=invoice_docx_path,
        apply_tax=apply_tax_var.get(),
        apply_card_fee=apply_card.get()
    )


    pdf_path = invoice_docx_path.replace(".docx", ".pdf")
    try:
        run(["open", pdf_path])  # Opens the converted PDF on macOS
    except Exception:
        pass

    # --- 5. Build base line items for QuickBooks ---

    def require_item_ref(name):
        ref = get_or_create_item(name, ACCESS_TOKEN)
        if not ref or not isinstance(ref, dict) or "value" not in ref:
            messagebox.showerror(
                "QuickBooks Error",
                f"Could not find or create the QuickBooks Item for: {name}\n\n"
                f"Log file: {LOG_PATH}"
                "If the problem persists, contact support:\n"
                "828-318-2202\nbenjaminzeidell@gmail.com"
            )
            return None
        return ref


    lines = []

    for item in invoice_items:
        unit_price = round(item["regular_price"], 2)
        qty = item["num_prints"]
        amount = round(unit_price * qty, 2)
    
        print_type = strip_emoji(item["print_type"]).strip()
        size = item.get("size", "").strip()
        title = item.get("title", item.get("linked_title", "Untitled")).strip()
    
        if print_type in bulk_pricing:
            description = f"{size} inches\n   {title}"
        else:
            description = f"{title}"
    
        item_ref = require_item_ref(print_type)
        if not item_ref:
            return
    
        lines.append({
            "DetailType": "SalesItemLineDetail",
            "Amount": amount,
            "Description": description,
            "SalesItemLineDetail": {
                "ItemRef": item_ref,
                "Qty": qty,
                "UnitPrice": unit_price,
                "TaxCodeRef": tax_code_ref,
            },
        })

        
    # --- 6. Separate discount lines ---
    
    # 6A. Volume Discount
    if volume_savings > 0:
        vol_item_ref = require_item_ref("Volume Discount")
        if not vol_item_ref:
            return
    
        vol_amount = round(volume_savings, 2)
        lines.append({
            "DetailType": "SalesItemLineDetail",
            "Amount": -vol_amount,
            "Description": "Volume Discount",
            "SalesItemLineDetail": {
                "ItemRef": vol_item_ref,
                "Qty": 1,
                "UnitPrice": -vol_amount,
                "TaxCodeRef": tax_code_ref,
            },
        })

    
    # 6B. Professional Discount
    if pro_savings > 0:
        pro_item_ref = require_item_ref("Professional Discount")
        if not pro_item_ref:
            return
            
        pro_amount = round(pro_savings, 2)
        lines.append({
            "DetailType": "SalesItemLineDetail",
            "Amount": -pro_amount,
            "Description": "Professional Discount",
            "SalesItemLineDetail": {
                "ItemRef": pro_item_ref,
                "Qty": 1,
                "UnitPrice": -pro_amount,
                "TaxCodeRef": tax_code_ref, 
                
            }
        })
    
    # 6C. Flat Discount ($)
    if flat_discount > 0:
        flat_item_ref = require_item_ref("Flat Discount")
        if not flat_item_ref:
            return
            
        flat_amount = round(flat_discount, 2)
        lines.append({
            "DetailType": "SalesItemLineDetail",
            "Amount": -flat_amount,
            "Description": "Flat Discount",
            "SalesItemLineDetail": {
                "ItemRef": flat_item_ref,
                "Qty": 1,
                "UnitPrice": -flat_amount,
                "TaxCodeRef": tax_code_ref,

            }
        })
    
    # 6D. Custom Percentage Discount
    if percent_discount_amt > 0:
        pct_item_ref = require_item_ref("Custom % Discount")
        if not pct_item_ref:
            return
            
        pct_amount = round(percent_discount_amt, 2)
        lines.append({
            "DetailType": "SalesItemLineDetail",
            "Amount": -pct_amount,
            "Description": f"Custom Discount ({summary.get('percent_discount', 0)}%)",
            "SalesItemLineDetail": {
                "ItemRef": pct_item_ref,
                "Qty": 1,
                "UnitPrice": -pct_amount,
                "TaxCodeRef": tax_code_ref,

            }
        })


    # --- 7. Card fee line ---

    if card_fee > 0:
        fee_item_ref = require_item_ref("Card Fee")
        if not fee_item_ref:
            return
            
        fee_amount = round(card_fee, 2)

        lines.append({
            "DetailType": "SalesItemLineDetail",
            "Amount": fee_amount,
            "Description": "Card Fee (3%)",
            "SalesItemLineDetail": {
                "ItemRef": fee_item_ref,
                "Qty": 1,
                "UnitPrice": fee_amount,
                "TaxCodeRef": tax_code_ref,

            },
        })

    # --- 8. Final invoice payload (including tax) ---

    invoice_data = {
        "CustomerRef": {"value": str(customer_id)},
        "Line": lines,
    }
    
    if apply_tax_var.get():
        invoice_data["TxnTaxDetail"] = {
            "TxnTaxCodeRef": {"value": "TAX"},
            "TotalTax": round(tax_amount, 2)
        }


    # --- 9. Send to QuickBooks ---

    try:
        response = requests.post(
            f"{BASE_URL}/v3/company/{REALM_ID}/invoice",
            headers={
                "Authorization": f"Bearer {ACCESS_TOKEN}",
                "Content-Type": "application/json",
                "Accept": "application/json",
            },
            json=invoice_data,
        )
    except requests.RequestException as e:
        logging.error("INVOICE_CREATE_REQUEST_EXCEPTION %s", {"error": str(e)})
        messagebox.showerror(
            "QuickBooks Connection Error",
            "Could not contact QuickBooks to create the invoice.\n\n"
            f"Error: {e}\n\nLog file: {LOG_PATH}"
            "If the problem persists, contact support:\n"
            "828-318-2202\nbenjaminzeidell@gmail.com"
        )
        return
    
    if response.status_code == 200:
        messagebox.showinfo("Success", "Invoice sent to QuickBooks!")
    else:
        log_qbo_error("invoice_create", response, extra={"customer_id": customer_id})
    
        tid = get_intuit_tid(response)
        messagebox.showerror(
            "QuickBooks Error",
            "QuickBooks rejected the invoice request.\n\n"
            f"Intuit TID: {tid}\n"
            f"Log file: {LOG_PATH}\n\n"
            "If the problem persists, contact support:\n"
            "828-318-2202\nbenjaminzeidell@gmail.com"
        )



# In[ ]:


import tkinter as tk
from tkinter import ttk, messagebox
import sys
import os
from collections import defaultdict

current_artist = ""
current_title = ""
draft_items = []
invoice_items = []
results_items = []
draft_titles = []
collapsed_titles = set()
color_index = 0
colors = ["#ccff00", "#00ccff", "#ff00cc", "#ffcc00", "#cc00ff", "#00ffcc", "#ff0000"]
custom_items_by_title = defaultdict(list)


# Bulk pricing data
bulk_pricing = {} # pricing data removed intentionally

def main_app():
    global root, draft_frame, invoice_frame, draft_box, invoice_box
    global height_entry, width_entry, num_prints_entry, original_height_entry, original_width_entry
    global results_inner_frame, results_canvas, results_box, apply_tax_var, apply_card
    global price_var_pro, price_var_regular
    global artist_first_entry, artist_last_entry, title_entry
    global custom_discount_var, custom_discount_dollar_var
    global custom_item_name_entry, custom_item_desc_entry
    global custom_item_qty_var, custom_item_price_var
    global custom_items_by_title

    if current_title and current_title not in draft_titles:
        draft_titles.append(current_title)

    """
    Integrated app with user inputs at the top, dynamic results in a scrollable grid layout, and a close button.
    """
    def calculate_results():
        ###### Not everything needs # of prints or print sizes, so these have to be conditional
        global current_title
        current_title = title_entry.get().strip()

        height_text = height_entry.get()
        width_text = width_entry.get()
        num_prints_text = num_prints_entry.get()
        
        height = float(height_text) if height_text.strip() else None
        width = float(width_text) if width_text.strip() else None
        num_prints = int(num_prints_text) if num_prints_text.strip() else None

        ##### Not everything needs original height, so they have to be conditional too
        original_height_text = original_height_entry.get()
        original_width_text = original_width_entry.get()
        
        original_height = float(original_height_text) if original_height_text.strip() else None
        original_width = float(original_width_text) if original_width_text.strip() else None

        """
        Create a block for each medium, including its title, itemized costs, and separator line.
        """
        
        # Check if at least one price checkbox is checked & clear results_inner_frame is selection is changed
        try:
            if not price_var_regular.get() and not price_var_pro.get():
                messagebox.showerror("⚠️ Pricing Error", "You must select at least one price option.")
                return  # Stop execution
            results_box.configure(state="normal")
            results_box.delete("1.0", tk.END)
            results_box.configure(state="disabled")
            results_items.clear()

            frame_costs = # pricing data removed intentionally

            gallery_stretching_fees = # pricing data removed intentinoally
            basic_stretching_fees = # pricing data removed intentionally

            bracer_bar_cost = # pricing data removed intentionally
            upcharge_72 = # pricing data removed intentionally
            
            color_index = 0

                        # Determine capture size category and cost
            capture_size = ""
            capture_price = 0.00
             
            if capture_var.get():  # ✅ If "Capture" is checked
                original_max_dim = max(original_height, original_width)
                if original_max_dim < 48:
                    capture_size = "Small"
                    capture_price = # pricing data removed intentionally
                elif 48 <= original_max_dim < 72:
                    capture_size = "Medium"
                    capture_price =# pricing data removed intentionally
                else:
                    capture_size = "Large"
                    capture_price = # pricing data removed intentionally

            # ✅ Always check for add-ons, even if print fields are filled
            addons_selected = (
                capture_var.get() or
                specialty_capture_var.get() or
                color_match_var.get() or
                monitor_match_var.get() or
                additional_rounds_var.get() > 0 or
                complex_wrap_var.get() or
                artist_first_entry.get().strip() or
                artist_last_entry.get().strip() or
                title_entry.get().strip()
            )

            
            if addons_selected:
                    # 🧹 Remove stale add-ons
                draft_data = {
                    "canvas_cost": 0.00,
                    "pro_canvas_cost": 0.00,
                    "print_type": "Add-On Only Order",
                    "size": "",
                    "total_cost": 0.00,
                    "pro_total_cost": 0.00,
                    "num_prints": 0,
                    "frame_cost": 0.00,
                    "stretching_fee": 0.00,
                    "bracer_cost": 0.00,
                    "upcharge": 0.00,
                    "color": "#FFFFFF",
                    "capture_price": capture_price if capture_var.get() else None,
                    "capture_size": f"{capture_size} Capture" if capture_var.get() else None,
                    "specialty_capture": "✨ Specialty Capture" if specialty_capture_var.get() else None,
                    "capture_var": capture_var,
                    "specialty_capture_var": specialty_capture_var,
                    "color_match_var": color_match_var,
                    "complex_wrap_var": complex_wrap_var,
                    "additional_rounds_var": additional_rounds_var,
                    "flashdrive_var": flashdrive_var,
                    "computer_time_var": computer_time_var,
                    "monitor_match_var": monitor_match_var,
                    "artist_first": artist_first_entry.get(),
                    "artist_last": artist_last_entry.get(),
                    "title": title_entry.get(),
                    "linked_title": current_title,
                    "original_height": original_height,
                    "original_width": original_width,
                    "custom_items": custom_items_by_title.get(current_title, []),
                    "calculate_results": calculate_results
                }

                send_to_draft(draft_data)
                update_draft_display()
                #print(f"✅ Adding item linked to title: {current_title}")


                
            if height is not None and width is not None and num_prints is not None:
                for print_type, prices in bulk_pricing.items():
                    if num_prints <= 4:
                        price_per_sqft = prices[0]
                        pro_price_per_sqft = prices[1]
                        pro_discount = (price_per_sqft - pro_price_per_sqft)
                        volume_discount = 0.00
                    elif num_prints <= 19:
                        price_per_sqft = prices[1]
                        pro_price_per_sqft = prices[2]
                        pro_discount = (price_per_sqft - pro_price_per_sqft)
                        volume_discount = (prices[1] - prices[0])
                    elif num_prints <= 49:
                        price_per_sqft = prices[2]
                        pro_price_per_sqft = prices[3]
                        pro_discount = (price_per_sqft - pro_price_per_sqft)
                        volume_discount = (prices[2] - prices[1])
                    elif num_prints <= 99:
                        price_per_sqft = prices[3]
                        pro_price_per_sqft = prices[4]
                        pro_discount = (price_per_sqft - pro_price_per_sqft)
                        volume_discount = (prices[3] - prices[2])
                    else:
                        price_per_sqft = prices[4]
                        pro_price_per_sqft = (prices[4] * (prices[4] / prices[3]))
                        pro_discount = (price_per_sqft - pro_price_per_sqft)
                        volume_discount = (prices[4] - prices[3])

                    #print(f"After print_type, prices in bulk_pricing.items() loop: {pro_discount=}")
                    #print(f"After print_type, prices in bulk_pricing.items() loop: {volume_discount=}")
    
                    gallery_wrap_adjustment = 3 if print_type.startswith("Canvas with T") else 0
                    adjusted_height = height + gallery_wrap_adjustment
                    adjusted_width = width + gallery_wrap_adjustment
                    area_in_sqft = (adjusted_height * adjusted_width) / 144
                    perimeter_in_feet = (2 * height + 2 * width) / 12
    
                    if print_type.startswith("Canvas with"):
                        frame_cost = perimeter_in_feet * frame_costs.get(print_type, frame_costs["Default"])
                        bracer_cost = 0
                        if max(height, width) >= 40:
                            bracer_cost = (min(height, width) / 12) * bracer_bar_cost
                        if max(height, width) >= 60:
                            bracer_cost *= 2
                    else:
                        frame_cost = 0
                        bracer_cost = 0
                        
                    if print_type.startswith("Canvas with"):
                        upcharge = upcharge_72 if max(height, width) >= 72 else 0
                    else:
                        upcharge = 0
    
                    # Determine the stretching fee dictionary based on the print type
                    if print_type.startswith("Canvas with T"):
                        stretching_fees = gallery_stretching_fees
                    elif print_type == "Canvas with Basic Stretch":
                        stretching_fees = basic_stretching_fees
                    else:
                        stretching_fees = {}
    
                    # Calculate the stretching fee
                    stretching_fee = 0
                    for category, (min_range, max_range, fee) in stretching_fees.items():
                        if min_range <= height + width <= max_range:
                            stretching_fee = fee
                            break
            
                    canvas_cost = area_in_sqft * price_per_sqft
                    pro_canvas_cost = area_in_sqft * pro_price_per_sqft
                    volume_discount_amt = num_prints * area_in_sqft * volume_discount
                    pro_discount_amt = num_prints * area_in_sqft * pro_discount
                    total_cost = canvas_cost + frame_cost + bracer_cost + upcharge + stretching_fee
                    pro_total_cost = pro_canvas_cost + frame_cost + bracer_cost + upcharge + stretching_fee

                    #print(f"After volume_discount_amt = num_prints * area_in_sqft * volume_discount assignment: {volume_discount_amt=}")
                    #print(f"After pro_discount_amt = num_prints * area_in_sqft * pro_discount assignment: {pro_discount_amt=}")
    
                    # Define a list of colors to cycle through
                    colors = ["#ccff00", "#00ccff", "#ff00cc", "#ffcc00", "#cc00ff", "#00ffcc", "#ff0000"]
                    
                    # Cycle through the colors
                    current_color = colors[color_index % len(colors)]

                    # Skip rendering if all prices are 0 and no prints
                    if (
                        (canvas_cost == 0 and pro_canvas_cost == 0) and
                        (frame_cost == 0 and stretching_fee == 0 and bracer_cost == 0 and upcharge == 0) and
                        (total_cost == 0 and pro_total_cost == 0)
                    ):
                        continue  # Skip to the next print_type

                    draft_data = {
                            "canvas_cost": canvas_cost if price_var_regular.get() else None,
                            "pro_canvas_cost": pro_canvas_cost if price_var_pro.get() else None,
                            "volume_discount_amt": volume_discount_amt,
                            "pro_discount_amt": pro_discount_amt if price_var_pro.get() else None,
                            "print_type": print_type,
                            "size": f"{height} x {width}",
                            "total_cost": total_cost if price_var_regular.get() else None,
                            "pro_total_cost": pro_total_cost if price_var_pro.get() else None,
                            "num_prints": num_prints,
                            "frame_cost": frame_cost,
                            "stretching_fee": stretching_fee,
                            "bracer_cost": bracer_cost,
                            "upcharge": upcharge,
                            "color": current_color,
                            "capture_price": capture_price if capture_var.get() else None,
                            "capture_size": f"{capture_size} Capture" if capture_var.get() else None,
                            "specialty_capture": "✨ Specialty Capture" if specialty_capture_var.get() else None,
                            "flashdrive_var": flashdrive_var,
                            "computer_time_var": computer_time_var,
                            "capture_var": capture_var,
                            "specialty_capture_var": specialty_capture_var,
                            "color_match_var": color_match_var,
                            "complex_wrap_var": complex_wrap_var,
                            "additional_rounds_var": additional_rounds_var,
                            "monitor_match_var": monitor_match_var,
                            "artist_first": artist_first_entry.get(),
                            "artist_last": artist_last_entry.get(),
                            "title": title_entry.get(),
                            "linked_title": current_title,
                            "original_height": original_height,
                            "original_width": original_width,
                            "custom_items": custom_items_by_title.get(current_title, []),
                            "calculate_results": calculate_results
                        }

                    draft_data["artist_first"] = artist_first_entry.get()
                    draft_data["artist_last"] = artist_last_entry.get()
                    draft_data["title"] = title_entry.get()
                    send_to_results(draft_data)

                    #print(f"After send_to_results(draft_data): {volume_discount_amt=}")
                    #print(f"After send_to_results(draft_data): {pro_discount_amt=}")
                    
                    color_index += 1

                    ##print(f"ARTIST = {artist_first_entry.get()} {artist_last_entry.get()}")
                    ##print(f"TITLE  = {title_entry.get()}")
                    #print(f"Draft items: {[item['print_type'] for item in draft_items]}")

        except ValueError:
            messagebox.showerror("⚠️ Invalid Input", "Height, Width, and Number of Prints cannot be blank.")

    def exit_program(event=None):
        """
        Fully terminates the application, ensuring no lingering taskbar icon.
        """
        global root
        try:
            if root and root.winfo_exists():
                root.quit()
                root.destroy()
        except Exception as e:
            pass
            #print(f"❌ Error closing the program: {e}")
        
        os._exit(0)  # ✅ Hard exit to completely terminate the process
    
    root = tk.Tk()
    root.withdraw()  # Hide the weird default root window

    app = tk.Toplevel(root)
    app.title("Print Cost Calculator")
    app.geometry("1500x1500")

    # Input frame
    input_frame = tk.Frame(app)
    input_frame.grid(row=0, column=0, columnspan=3, sticky="ew", padx=10, pady=10)
    
    # Row 0 - Artist's Name
    tk.Label(input_frame, text="Artist's First Name:", font=("Avenir Next", 12)).grid(row=0, column=0, padx=5, pady=5, sticky="e")
    artist_first_entry = tk.Entry(input_frame, font=("Avenir Next", 12))
    artist_first_entry.grid(row=0, column=1, padx=5, pady=5)
    
    tk.Label(input_frame, text="Artist's Last Name:", font=("Avenir Next", 12)).grid(row=0, column=2, padx=5, pady=5, sticky="e")
    artist_last_entry = tk.Entry(input_frame, font=("Avenir Next", 12))
    artist_last_entry.grid(row=0, column=3, padx=5, pady=5)

    custom_discount_var = tk.DoubleVar(value=0.0) # Doublevar for when it's a float
    tk.Label(input_frame, text="💯                   Discount (%)", font=("Avenir Next", 12)).grid(row=0, column=4, padx=1, pady=5, sticky="e")
    tk.Spinbox(input_frame, from_=0.0, to=100, increment=0.01, font=("Avenir Next", 12), width=0, textvariable=custom_discount_var).grid(row=0, column=4, padx=50, pady=5, sticky="w")

    custom_discount_dollar_var = tk.DoubleVar(value=0.0) # Doublevar for when it's a float
    tk.Label(input_frame, text="💰                    Discount ($)", font=("Avenir Next", 12)).grid(row=0, column=5, padx=1, pady=5, sticky="e")
    tk.Spinbox(input_frame, from_=0.0, to=100, increment=0.01, font=("Avenir Next", 12), width=0, textvariable=custom_discount_dollar_var).grid(row=0, column=5, padx=35, pady=5, sticky="w")

    # Row 1 - Pricing Checkboxes & early addons
    price_var_regular = tk.BooleanVar(value=True)
    price_var_pro = tk.BooleanVar(value=False)
    apply_tax_var = tk.BooleanVar(value=True)  # Value = T/F manages behavior on app-open
    apply_card = tk.BooleanVar(value=True)
    additional_rounds_var = tk.IntVar(value=0)  # Intvar for when it's always an int
    computer_time_var = tk.DoubleVar(value=0.0) # Doublevar for when it's a float
    flashdrive_var = tk.IntVar(value=0)
    

    tk.Checkbutton(input_frame, text="Regular Price", font=("Avenir Next", 12), variable=price_var_regular).grid(row=1, column=0, padx=5, pady=5, sticky="w")
    tk.Checkbutton(input_frame, text="Professional Price", font=("Avenir Next", 12), variable=price_var_pro).grid(row=1, column=1, padx=5, pady=5, sticky="w")
    tk.Checkbutton(input_frame, text="Sales Tax", font=("Avenir Next", 12), variable=apply_tax_var).grid(row=1, column=2, padx=5, pady=5, sticky="w")
    tk.Checkbutton(input_frame, text="Card Fee", font=("Avenir Next", 12), variable=apply_card).grid(row=1, column=3, padx=0, pady=5, sticky="w")
    
    tk.Label(input_frame, text="🕖                   Computer Time", font=("Avenir Next", 12)).grid(row=1, column=4, padx=1, pady=5, sticky="e")
    tk.Spinbox(input_frame, from_=0.0, to=100, increment=0.25, font=("Avenir Next", 12), width=0, textvariable=computer_time_var).grid(row=1, column=4, padx=35, pady=5, sticky="w")

    tk.Label(input_frame, text="💿                   Flashdrive", font=("Avenir Next", 12)).grid(row=1, column=5, padx=1, pady=5)
    tk.Spinbox(input_frame, from_=0, to=100, font=("Avenir Next", 12), width=0, textvariable=flashdrive_var).grid(row=1, column=5, padx=35, pady=5, sticky="w")
    
    ttk.Separator(input_frame, orient="vertical").grid(row=0, column=7, rowspan=7, sticky="ns", padx=(10, 0))
    ttk.Separator(input_frame, orient="horizontal").grid(row=2, column=0, columnspan=7, sticky="ew", pady=(5, 5))
    
    # Row 2 - Add-ons
    capture_var = tk.BooleanVar()
    specialty_capture_var = tk.BooleanVar()
    color_match_var = tk.BooleanVar()
    monitor_match_var = tk.BooleanVar()
    complex_wrap_var = tk.BooleanVar()

    custom_item_qty_var = tk.IntVar(value=0)
    custom_item_price_var = tk.DoubleVar(value=0.00)

    
    tk.Checkbutton(input_frame, text="📸 Capture", font=("Avenir Next", 12), variable=capture_var).grid(row=3, column=0, padx=5, pady=5, sticky="w")
    tk.Checkbutton(input_frame, text="✨ Specialty Capture", font=("Avenir Next", 12), variable=specialty_capture_var).grid(row=3, column=1, padx=5, pady=5, sticky="w")
    tk.Checkbutton(input_frame, text="🎨 Basic Color Match", font=("Avenir Next", 12), variable=color_match_var).grid(row=3, column=2, padx=5, pady=5, sticky="w")
    tk.Checkbutton(input_frame, text="🖥️ Monitor Match", font=("Avenir Next", 12), variable=monitor_match_var).grid(row=3, column=3, padx=5, pady=5, sticky="w")
    tk.Checkbutton(input_frame, text="🐩 Complex Image Wrap", font=("Avenir Next", 12), variable=complex_wrap_var).grid(row=3, column=4, padx=5, pady=5, sticky="w")
    tk.Label(input_frame, text="💻             Additional Rounds", font=("Avenir Next", 12)).grid(row=3, column=5, padx=1, pady=5, sticky="e")
    tk.Spinbox(input_frame, from_=0, to=100, font=("Avenir Next", 12), width=0, textvariable=additional_rounds_var).grid(row=3, column=5, padx=25, pady=5, sticky="w")

    # Row 3 - Original dims & title
    tk.Label(input_frame, text="Original Height (in):", font=("Avenir Next", 12)).grid(row=4, column=0, padx=5, pady=5, sticky="e")
    original_height_entry = tk.Entry(input_frame, font=("Avenir Next", 12))
    original_height_entry.grid(row=4, column=1, padx=5, pady=5)
    
    tk.Label(input_frame, text="Original Width (in):", font=("Avenir Next", 12)).grid(row=4, column=2, padx=5, pady=5, sticky="e")
    original_width_entry = tk.Entry(input_frame, font=("Avenir Next", 12))
    original_width_entry.grid(row=4, column=3, padx=5, pady=5)
    
    tk.Label(input_frame, text="Title of Piece:", font=("Avenir Next", 12)).grid(row=4, column=4, padx=5, pady=5, sticky="e")
    title_entry = tk.Entry(input_frame, font=("Avenir Next", 12))
    title_entry.grid(row=4, column=5, padx=5, pady=5)

    artist_first_entry.bind("<KeyRelease>", on_artist_or_title_change)
    artist_last_entry.bind("<KeyRelease>", on_artist_or_title_change)
    title_entry.bind("<KeyRelease>", on_artist_or_title_change)

    ttk.Separator(input_frame, orient="horizontal").grid(row=5, column=0, columnspan=7, sticky="ew", pady=(5, 5))


    # Row 4 - Print details & calculate
    tk.Label(input_frame, text="Number of Prints:", font=("Avenir Next", 12)).grid(row=6, column=0, padx=5, pady=5, sticky="e")
    num_prints_entry = tk.Entry(input_frame, font=("Avenir Next", 12))
    num_prints_entry.grid(row=6, column=1, padx=5, pady=5)
    
    tk.Label(input_frame, text="Print Height (in):", font=("Avenir Next", 12)).grid(row=6, column=2, padx=5, pady=5, sticky="e")
    height_entry = tk.Entry(input_frame, font=("Avenir Next", 12))
    height_entry.grid(row=6, column=3, padx=5, pady=5)
    
    tk.Label(input_frame, text="Print Width (in):", font=("Avenir Next", 12)).grid(row=6, column=4, padx=5, pady=5, sticky="e")
    width_entry = tk.Entry(input_frame, font=("Avenir Next", 12))
    width_entry.grid(row=6, column=5, padx=5, pady=5)
    
    tk.Button(input_frame, text="Calculate", font=("Avenir Next", 12, "bold"), command=calculate_results).grid(row=6, column=6, padx=10, pady=10)
    
    ttk.Separator(input_frame, orient="horizontal").grid(row=7, column=0, columnspan=7, sticky="ew", pady=(5, 5))


    # Row 5 (custom item name & desc)
    tk.Label(input_frame, text="                                                  🆕 Custom Item Name", font=("Avenir Next", 12)).grid(row=0, column=8, padx=5, pady=5, sticky="w")
    custom_item_name_entry = tk.Entry(input_frame, font=("Avenir Next", 12), width=20)
    custom_item_name_entry.grid(row=0, column=8, padx=5, pady=5, sticky="w")
    
    tk.Label(input_frame, text="                                                  📝 Description", font=("Avenir Next", 12)).grid(row=1, column=8, padx=5, pady=5, sticky="w")
    custom_item_desc_entry = tk.Entry(input_frame, font=("Avenir Next", 12), width=20)
    custom_item_desc_entry.grid(row=1, column=8, padx=5, pady=5, sticky="w")
    
    # Row 7 (quantity & price)
    tk.Label(input_frame, text="                    Qty", font=("Avenir Next", 12)).grid(row=2, column=8, padx=5, pady=5, sticky="w")
    tk.Spinbox(input_frame, from_=0, to=999, font=("Avenir Next", 12), width=5, textvariable=custom_item_qty_var).grid(row=2, column=8, padx=5, pady=5, sticky="w")
    
    tk.Label(input_frame, text="                           Price", font=("Avenir Next", 12)).grid(row=2, column=8, padx=90, pady=5, sticky="w")
    tk.Spinbox(input_frame, from_=0.0, to=9999.0, increment=0.01, font=("Avenir Next", 12), width=8, textvariable=custom_item_price_var).grid(row=2, column=8, padx=90, pady=5, sticky="w")

    custom_item_add = tk.Button(
        input_frame,
        text="Add Custom Item",
        font=("Avenir Next", 12, "bold"),
        command=add_custom_item_to_draft,
        bg="#222222",
        bd=0, highlightthickness=0, padx=2, pady=0)
    
    custom_item_add.grid(row=3, column=8, padx=5, pady=5, sticky="w")

    
    # Make sure the root window expands properly
    app.grid_rowconfigure(1, weight=1)  # Allows main content to expand
    app.grid_columnconfigure(0, weight=1)  # Left column
    app.grid_columnconfigure(1, weight=1)  # Center column (Draft Box)
    app.grid_columnconfigure(2, weight=1)  # Right column (Final Invoice)
    
    results_frame = tk.Frame(app)
    results_frame.grid(row=1, column=0, sticky="nsew", padx=0, pady=0)  
    
    results_frame.grid_rowconfigure(0, weight=1)
    results_frame.grid_columnconfigure(0, weight=1)
    
    results_box = tk.Text(results_frame, height=15, width=50, font=("Monaco", 12), bg="#222222", fg="white", state="disabled")
    results_box.grid(row=0, column=0, sticky="nsew")
    
    scrollbar_y = ttk.Scrollbar(results_frame, orient="vertical", command=results_box.yview)
    scrollbar_y.grid(row=0, column=1, sticky="ns")
    results_box.config(yscrollcommand=scrollbar_y.set)

    draft_frame = tk.Frame(app, bg="#333333")
    draft_frame.grid(row=1, column=1, sticky="nsew", padx=0, pady=0)  
    
    invoice_frame = tk.Frame(app, bg="#333333")
    invoice_frame.grid(row=1, column=2, sticky="nsew", padx=0, pady=0)

    #draft_scrollbar = ttk.Scrollbar(draft_frame, orient="vertical", command=lambda *args: draft_box.yview(*args))
    draft_box = tk.Text(draft_frame, height=15, width=50, font=("Monaco", 12), bg="#222222", fg="white", state="disabled")
    draft_box.grid(row=0, column=0, sticky="nsew")
    #draft_scrollbar.grid(row=0, column=1, sticky="ns")

    draft_scrollbar = ttk.Scrollbar(draft_frame, orient="vertical", command=draft_box.yview)
    draft_scrollbar.grid(row=0, column=1, sticky="ns")
    draft_box.config(yscrollcommand=draft_scrollbar.set)


    draft_frame.grid_rowconfigure(0, weight=1)
    draft_frame.grid_columnconfigure(0, weight=1)

    draft_button_row = tk.Frame(draft_frame, bg="#222222")
    draft_button_row.grid(row=1, column=0, columnspan=2, pady=5)
    
    send_to_invoice_button = tk.Button(draft_button_row, text="Send to Draft", font=("Avenir Next", 12, "bold"), command=send_to_invoice)
    send_to_invoice_button.pack(side="left", padx=(0, 10))
    
    clear_draft_button = tk.Button(draft_button_row, text="Clear All", font=("Avenir Next", 12, "bold"), command=clear_draft)
    clear_draft_button.pack(side="left")
    
    invoice_box = tk.Text(invoice_frame, height=15, width=50, font=("Monaco", 12), bg="#222222", fg="white", state="disabled")
    invoice_box.grid(row=0, column=0, sticky="nsew")

    invoice_scrollbar = ttk.Scrollbar(invoice_frame, orient="vertical", command=invoice_box.yview)
    invoice_scrollbar.grid(row=0, column=1, sticky="ns")
    invoice_box.config(yscrollcommand=invoice_scrollbar.set)

    #invoice_scrollbar = ttk.Scrollbar(invoice_frame, orient="vertical", command=lambda *args: invoice_box.yview(*args))
    invoice_frame.grid_rowconfigure(0, weight=1)
    invoice_frame.grid_columnconfigure(0, weight=1)
    #invoice_scrollbar.grid(row=0, column=1, sticky="ns")
    
    invoice_button_row = tk.Frame(invoice_frame, bg="#222222")
    invoice_button_row.grid(row=1, column=0, columnspan=2, pady=5)
    
    send_to_quickbooks_button = tk.Button(invoice_button_row, text="Generate Final Invoice", font=("Avenir Next", 12, "bold"), command=send_to_quickbooks)
    send_to_quickbooks_button.pack(side="left", padx=(0, 10))
    
    clear_invoice_button = tk.Button(invoice_button_row, text="Clear All", font=("Avenir Next", 12, "bold"), command=clear_invoice)
    clear_invoice_button.pack(side="left")
    
    # Bind scrolling to the canvas
    results_box.bind("<Enter>", lambda e: results_box.bind_all("<MouseWheel>", scroll_mac))
    results_box.bind("<Leave>", lambda e: results_box.unbind_all("<MouseWheel>"))

    close_button = tk.Button(app, text="Close", font=("Avenir Next", 12, "bold"), command=exit_program)
    close_button.grid(row=2, column=1, pady=10)

    # Bind Enter key to submit inputs
    app.bind("<Return>", lambda event: calculate_results())

    # Bind click in draft box to focus on nearest title AND re-trigger results
    draft_box.bind("<Button-1>", lambda e: handle_click_in_draft_box(e, calculate_results))


    # Bind Escape key to exit program
    app.bind("<Escape>", exit_program)

    # Bind Up and Down arrow keys to scroll functionality
    results_box.bind("<Up>", lambda e: results_box.yview_scroll(-1, "units"))
    results_box.bind("<Down>", lambda e: results_box.yview_scroll(1, "units"))

    def scroll_mac(widget):
        def on_mousewheel(event):
            widget.yview_scroll(-1 * event.delta, "units")
            return "break"
        return on_mousewheel
    
    results_box.bind("<MouseWheel>", scroll_mac(results_box))
    draft_box.bind("<MouseWheel>", scroll_mac(draft_box))
    invoice_box.bind("<MouseWheel>", scroll_mac(invoice_box))

    def bind_smooth_scroll(box):
        def on_mousewheel(event):
            box.yview_scroll(-1 * event.delta, "units")
            return "break"
        box.bind("<MouseWheel>", on_mousewheel)

    bind_smooth_scroll(results_box)
    bind_smooth_scroll(draft_box)
    bind_smooth_scroll(invoice_box)

    app.bind_all("<MouseWheel>", scroll_mac(results_box))

    artist_first_entry.focus()

    root.mainloop()

# Run the main app
main_app()